import PyPDF2
import pdfplumber
import docx

def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ''

    # Extract tables using pdfplumber
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    text += ' | '.join(str(cell) if cell else '' for cell in row) + '\n'

    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    # Extract tables from DOCX
    for table in doc.tables:
        for row in table.rows:
            text += ' | '.join(cell.text.strip() for cell in row.cells) + '\n'

    return text


def extract_text_from_resume(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")